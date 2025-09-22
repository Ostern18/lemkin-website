"""
Core OCR and document processing functionality for the Lemkin OCR Toolkit.

This module provides comprehensive document processing capabilities including:
- Multi-engine OCR with confidence scoring
- Layout analysis and structure preservation
- Text extraction with formatting retention
- Multi-language support and character encoding handling
- Legal document-specific processing optimizations
"""

import logging
import tempfile
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from uuid import uuid4

import cv2
import numpy as np
import pandas as pd
import pytesseract
import torch
from PIL import Image
from pydantic import BaseModel, Field, validator
from transformers import pipeline

# Document processing imports
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
from pdf2image import convert_from_path

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types for processing."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    IMAGE = "image"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"


class OCREngine(str, Enum):
    """Available OCR engines."""
    TESSERACT = "tesseract"
    EASYOCR = "easyocr"
    PADDLEOCR = "paddleocr"
    AZURE = "azure"
    GOOGLE = "google"
    AWS = "aws"


class LanguageCode(str, Enum):
    """Supported language codes for OCR."""
    EN = "eng"  # English
    ES = "spa"  # Spanish
    FR = "fra"  # French
    DE = "deu"  # German
    IT = "ita"  # Italian
    PT = "por"  # Portuguese
    RU = "rus"  # Russian
    ZH = "chi_sim"  # Chinese Simplified
    JA = "jpn"  # Japanese
    AR = "ara"  # Arabic
    AUTO = "auto"  # Auto-detect


class ConfidenceLevel(str, Enum):
    """OCR confidence levels."""
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    VERY_LOW = "very_low"


class TextBlock(BaseModel):
    """Represents a block of extracted text with metadata."""
    text: str = Field(description="Extracted text content")
    confidence: float = Field(ge=0.0, le=1.0, description="OCR confidence score")
    bbox: Tuple[int, int, int, int] = Field(description="Bounding box (x, y, width, height)")
    page_number: int = Field(ge=1, description="Page number (1-indexed)")
    block_type: str = Field(default="text", description="Type of text block")
    language: Optional[LanguageCode] = Field(default=None, description="Detected language")
    font_info: Optional[Dict[str, Any]] = Field(default=None, description="Font information")


class DocumentMetadata(BaseModel):
    """Document metadata extracted during processing."""
    title: Optional[str] = Field(default=None, description="Document title")
    author: Optional[str] = Field(default=None, description="Document author")
    subject: Optional[str] = Field(default=None, description="Document subject")
    creator: Optional[str] = Field(default=None, description="Creator application")
    producer: Optional[str] = Field(default=None, description="Producer application")
    creation_date: Optional[datetime] = Field(default=None, description="Creation date")
    modification_date: Optional[datetime] = Field(default=None, description="Last modification date")
    page_count: int = Field(ge=0, description="Total number of pages")
    language: Optional[str] = Field(default=None, description="Document language")
    security_settings: Optional[Dict[str, Any]] = Field(default=None, description="Security settings")


class LayoutElement(BaseModel):
    """Represents a layout element (table, image, heading, etc.)."""
    element_type: str = Field(description="Type of layout element")
    bbox: Tuple[int, int, int, int] = Field(description="Bounding box")
    page_number: int = Field(ge=1, description="Page number")
    content: Optional[str] = Field(default=None, description="Text content if applicable")
    attributes: Dict[str, Any] = Field(default_factory=dict, description="Element attributes")
    confidence: float = Field(ge=0.0, le=1.0, default=1.0, description="Detection confidence")


class LayoutAnalysis(BaseModel):
    """Complete layout analysis of a document."""
    analysis_id: str = Field(default_factory=lambda: str(uuid4()))
    document_path: str = Field(description="Path to analyzed document")
    elements: List[LayoutElement] = Field(description="Detected layout elements")
    page_layouts: Dict[int, List[LayoutElement]] = Field(description="Elements grouped by page")
    reading_order: List[str] = Field(description="Suggested reading order of elements")
    analysis_timestamp: datetime = Field(default_factory=datetime.utcnow)


class OCRResult(BaseModel):
    """Result of OCR processing on a document."""
    ocr_id: str = Field(default_factory=lambda: str(uuid4()))
    document_path: str = Field(description="Path to processed document")
    engine_used: OCREngine = Field(description="OCR engine used")
    text_blocks: List[TextBlock] = Field(description="Extracted text blocks")
    full_text: str = Field(description="Complete extracted text")
    total_confidence: float = Field(ge=0.0, le=1.0, description="Overall confidence score")
    languages_detected: List[LanguageCode] = Field(description="Detected languages")
    processing_time: float = Field(description="Processing time in seconds")
    pages_processed: int = Field(ge=0, description="Number of pages processed")
    ocr_timestamp: datetime = Field(default_factory=datetime.utcnow)


class ExtractionResult(BaseModel):
    """Result of text extraction from various document formats."""
    extraction_id: str = Field(default_factory=lambda: str(uuid4()))
    document_path: str = Field(description="Path to processed document")
    document_type: DocumentType = Field(description="Type of document processed")
    extracted_text: str = Field(description="Extracted text content")
    metadata: Optional[DocumentMetadata] = Field(default=None, description="Document metadata")
    tables: List[pd.DataFrame] = Field(default_factory=list, description="Extracted tables")
    images: List[Dict[str, Any]] = Field(default_factory=list, description="Extracted images")
    extraction_method: str = Field(description="Method used for extraction")
    success: bool = Field(description="Whether extraction was successful")
    processing_time: float = Field(description="Processing time in seconds")
    extraction_timestamp: datetime = Field(default_factory=datetime.utcnow)


class DocumentAnalysis(BaseModel):
    """Comprehensive document analysis combining OCR, layout, and extraction."""
    analysis_id: str = Field(default_factory=lambda: str(uuid4()))
    document_path: str = Field(description="Path to analyzed document")
    document_type: DocumentType = Field(description="Document type")
    ocr_result: Optional[OCRResult] = Field(default=None, description="OCR results")
    layout_analysis: Optional[LayoutAnalysis] = Field(default=None, description="Layout analysis")
    extraction_result: Optional[ExtractionResult] = Field(default=None, description="Text extraction results")
    final_text: str = Field(description="Final processed text")
    quality_score: float = Field(ge=0.0, le=1.0, description="Overall processing quality score")
    recommendations: List[str] = Field(default_factory=list, description="Processing recommendations")
    analysis_timestamp: datetime = Field(default_factory=datetime.utcnow)


class DocumentProcessor:
    """Main document processor supporting multiple formats and OCR engines."""

    def __init__(self, default_engine: OCREngine = OCREngine.TESSERACT):
        """Initialize document processor.

        Args:
            default_engine: Default OCR engine to use
        """
        self.default_engine = default_engine
        self.ocr_engine = OCREngine()
        self.layout_analyzer = LayoutAnalyzer()
        self.text_extractor = TextExtractor()

        logger.info(f"Initialized DocumentProcessor with engine {default_engine}")

    def process_document(
        self,
        document_path: Path,
        ocr_engine: Optional[OCREngine] = None,
        languages: Optional[List[LanguageCode]] = None,
        include_layout: bool = True,
        include_ocr: bool = True,
        include_extraction: bool = True
    ) -> DocumentAnalysis:
        """Process document with comprehensive analysis.

        Args:
            document_path: Path to document file
            ocr_engine: OCR engine to use (None for default)
            languages: Target languages for OCR
            include_layout: Whether to perform layout analysis
            include_ocr: Whether to perform OCR
            include_extraction: Whether to perform text extraction

        Returns:
            DocumentAnalysis with complete results
        """
        start_time = datetime.utcnow()

        try:
            logger.info(f"Processing document: {document_path}")

            # Determine document type
            doc_type = self._determine_document_type(document_path)

            analysis = DocumentAnalysis(
                document_path=str(document_path),
                document_type=doc_type,
                final_text=""
            )

            # Perform text extraction for native formats
            if include_extraction and doc_type in [DocumentType.PDF, DocumentType.DOCX, DocumentType.XLSX]:
                analysis.extraction_result = self.text_extractor.extract_text(
                    document_path, doc_type
                )

            # Perform OCR for images and scanned documents
            if include_ocr and (doc_type == DocumentType.IMAGE or self._requires_ocr(document_path, doc_type)):
                analysis.ocr_result = self.ocr_engine.perform_ocr(
                    document_path, ocr_engine or self.default_engine, languages
                )

            # Perform layout analysis
            if include_layout:
                analysis.layout_analysis = self.layout_analyzer.analyze_layout(document_path)

            # Combine results to create final text
            analysis.final_text = self._combine_results(analysis)
            analysis.quality_score = self._calculate_quality_score(analysis)
            analysis.recommendations = self._generate_recommendations(analysis)

            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.info(f"Document processing completed in {processing_time:.2f}s")

            return analysis

        except Exception as e:
            logger.error(f"Document processing failed for {document_path}: {e}")
            raise

    def _determine_document_type(self, document_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        ext = document_path.suffix.lower()

        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLS,
            '.txt': DocumentType.TXT,
            '.rtf': DocumentType.RTF,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.tiff': DocumentType.IMAGE,
            '.bmp': DocumentType.IMAGE,
            '.gif': DocumentType.IMAGE,
        }

        return type_mapping.get(ext, DocumentType.IMAGE)

    def _requires_ocr(self, document_path: Path, doc_type: DocumentType) -> bool:
        """Determine if document requires OCR processing."""
        if doc_type == DocumentType.IMAGE:
            return True

        # Check if PDF is scanned (simplified heuristic)
        if doc_type == DocumentType.PDF:
            try:
                doc = fitz.open(str(document_path))
                for page in doc:
                    text = page.get_text()
                    if len(text.strip()) < 50:  # Very little text suggests scanned PDF
                        return True
                return False
            except Exception:
                return True

        return False

    def _combine_results(self, analysis: DocumentAnalysis) -> str:
        """Combine results from different processing methods."""
        texts = []

        # Prioritize extraction results for native formats
        if analysis.extraction_result and analysis.extraction_result.success:
            texts.append(analysis.extraction_result.extracted_text)

        # Add OCR results if available and significant
        if analysis.ocr_result and analysis.ocr_result.full_text:
            # Only add OCR if extraction failed or text is substantially different
            if not texts or len(analysis.ocr_result.full_text) > len(texts[0]) * 1.5:
                texts.append(analysis.ocr_result.full_text)

        return "\n\n".join(texts) if texts else ""

    def _calculate_quality_score(self, analysis: DocumentAnalysis) -> float:
        """Calculate overall quality score for the analysis."""
        scores = []

        if analysis.ocr_result:
            scores.append(analysis.ocr_result.total_confidence)

        if analysis.extraction_result:
            scores.append(1.0 if analysis.extraction_result.success else 0.0)

        return sum(scores) / len(scores) if scores else 0.0

    def _generate_recommendations(self, analysis: DocumentAnalysis) -> List[str]:
        """Generate processing recommendations based on results."""
        recommendations = []

        if analysis.ocr_result and analysis.ocr_result.total_confidence < 0.7:
            recommendations.append("OCR confidence is low - consider preprocessing image or using different engine")

        if analysis.quality_score < 0.8:
            recommendations.append("Overall quality is below optimal - review results carefully")

        if not analysis.final_text.strip():
            recommendations.append("No text was extracted - verify document format and content")

        return recommendations


class OCREngine:
    """OCR engine wrapper supporting multiple OCR backends."""

    def __init__(self):
        """Initialize OCR engine."""
        self.engines = {
            OCREngine.TESSERACT: self._tesseract_ocr,
            OCREngine.EASYOCR: self._easyocr_ocr,
            OCREngine.PADDLEOCR: self._paddleocr_ocr,
        }
        logger.info("Initialized OCR engine with multiple backends")

    def perform_ocr(
        self,
        document_path: Path,
        engine: OCREngine = OCREngine.TESSERACT,
        languages: Optional[List[LanguageCode]] = None
    ) -> OCRResult:
        """Perform OCR on document using specified engine.

        Args:
            document_path: Path to document or image file
            engine: OCR engine to use
            languages: Target languages for OCR

        Returns:
            OCRResult with extracted text and metadata
        """
        start_time = datetime.utcnow()

        try:
            logger.info(f"Starting OCR with {engine} on {document_path}")

            # Convert document to images if necessary
            images = self._prepare_images(document_path)

            # Perform OCR using selected engine
            if engine not in self.engines:
                raise ValueError(f"Unsupported OCR engine: {engine}")

            ocr_func = self.engines[engine]
            text_blocks = []
            languages_detected = set()

            for page_num, image in enumerate(images, 1):
                page_blocks = ocr_func(image, languages, page_num)
                text_blocks.extend(page_blocks)

                # Collect detected languages
                for block in page_blocks:
                    if block.language:
                        languages_detected.add(block.language)

            # Combine all text
            full_text = "\n".join([block.text for block in text_blocks if block.text.strip()])

            # Calculate overall confidence
            if text_blocks:
                total_confidence = sum(block.confidence for block in text_blocks) / len(text_blocks)
            else:
                total_confidence = 0.0

            processing_time = (datetime.utcnow() - start_time).total_seconds()

            result = OCRResult(
                document_path=str(document_path),
                engine_used=engine,
                text_blocks=text_blocks,
                full_text=full_text,
                total_confidence=total_confidence,
                languages_detected=list(languages_detected),
                processing_time=processing_time,
                pages_processed=len(images)
            )

            logger.info(f"OCR completed: {len(text_blocks)} blocks, {total_confidence:.2f} confidence")
            return result

        except Exception as e:
            logger.error(f"OCR failed for {document_path}: {e}")
            raise

    def _prepare_images(self, document_path: Path) -> List[Image.Image]:
        """Convert document to images for OCR processing."""
        doc_type = document_path.suffix.lower()

        if doc_type == '.pdf':
            # Convert PDF to images
            return convert_from_path(str(document_path), dpi=300)
        elif doc_type in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            # Load single image
            return [Image.open(document_path)]
        else:
            raise ValueError(f"Unsupported document type for OCR: {doc_type}")

    def _tesseract_ocr(
        self,
        image: Image.Image,
        languages: Optional[List[LanguageCode]],
        page_number: int
    ) -> List[TextBlock]:
        """Perform OCR using Tesseract engine."""
        try:
            # Configure language
            lang_string = "+".join([lang.value for lang in languages]) if languages else "eng"

            # Perform OCR with detailed output
            data = pytesseract.image_to_data(
                image,
                lang=lang_string,
                output_type=pytesseract.Output.DICT,
                config='--psm 6'  # Uniform block of text
            )

            text_blocks = []
            for i in range(len(data['text'])):
                text = data['text'][i].strip()
                if text and int(data['conf'][i]) > 0:  # Only keep text with confidence > 0
                    block = TextBlock(
                        text=text,
                        confidence=int(data['conf'][i]) / 100.0,  # Convert to 0-1 scale
                        bbox=(
                            int(data['left'][i]),
                            int(data['top'][i]),
                            int(data['width'][i]),
                            int(data['height'][i])
                        ),
                        page_number=page_number,
                        block_type="text"
                    )
                    text_blocks.append(block)

            return text_blocks

        except Exception as e:
            logger.error(f"Tesseract OCR failed: {e}")
            return []

    def _easyocr_ocr(
        self,
        image: Image.Image,
        languages: Optional[List[LanguageCode]],
        page_number: int
    ) -> List[TextBlock]:
        """Perform OCR using EasyOCR engine."""
        try:
            import easyocr

            # Map language codes
            lang_codes = ['en'] if not languages else [self._map_to_easyocr_lang(lang) for lang in languages]

            reader = easyocr.Reader(lang_codes)
            results = reader.readtext(np.array(image))

            text_blocks = []
            for bbox, text, confidence in results:
                if text.strip():
                    # Convert bbox format
                    x_coords = [point[0] for point in bbox]
                    y_coords = [point[1] for point in bbox]
                    x, y = int(min(x_coords)), int(min(y_coords))
                    w, h = int(max(x_coords) - x), int(max(y_coords) - y)

                    block = TextBlock(
                        text=text,
                        confidence=confidence,
                        bbox=(x, y, w, h),
                        page_number=page_number,
                        block_type="text"
                    )
                    text_blocks.append(block)

            return text_blocks

        except Exception as e:
            logger.error(f"EasyOCR failed: {e}")
            return []

    def _paddleocr_ocr(
        self,
        image: Image.Image,
        languages: Optional[List[LanguageCode]],
        page_number: int
    ) -> List[TextBlock]:
        """Perform OCR using PaddleOCR engine."""
        try:
            from paddleocr import PaddleOCR

            # Initialize PaddleOCR
            lang = 'en' if not languages else self._map_to_paddle_lang(languages[0])
            ocr = PaddleOCR(use_angle_cls=True, lang=lang)

            results = ocr.ocr(np.array(image), cls=True)

            text_blocks = []
            if results and results[0]:
                for line in results[0]:
                    bbox, (text, confidence) = line
                    if text.strip():
                        # Convert bbox format
                        x_coords = [point[0] for point in bbox]
                        y_coords = [point[1] for point in bbox]
                        x, y = int(min(x_coords)), int(min(y_coords))
                        w, h = int(max(x_coords) - x), int(max(y_coords) - y)

                        block = TextBlock(
                            text=text,
                            confidence=confidence,
                            bbox=(x, y, w, h),
                            page_number=page_number,
                            block_type="text"
                        )
                        text_blocks.append(block)

            return text_blocks

        except Exception as e:
            logger.error(f"PaddleOCR failed: {e}")
            return []

    def _map_to_easyocr_lang(self, lang: LanguageCode) -> str:
        """Map LanguageCode to EasyOCR language code."""
        mapping = {
            LanguageCode.EN: 'en',
            LanguageCode.ES: 'es',
            LanguageCode.FR: 'fr',
            LanguageCode.DE: 'de',
            LanguageCode.IT: 'it',
            LanguageCode.PT: 'pt',
            LanguageCode.RU: 'ru',
            LanguageCode.ZH: 'ch_sim',
            LanguageCode.JA: 'ja',
            LanguageCode.AR: 'ar',
        }
        return mapping.get(lang, 'en')

    def _map_to_paddle_lang(self, lang: LanguageCode) -> str:
        """Map LanguageCode to PaddleOCR language code."""
        mapping = {
            LanguageCode.EN: 'en',
            LanguageCode.ES: 'es',
            LanguageCode.FR: 'fr',
            LanguageCode.DE: 'german',
            LanguageCode.IT: 'it',
            LanguageCode.PT: 'pt',
            LanguageCode.RU: 'ru',
            LanguageCode.ZH: 'ch',
            LanguageCode.JA: 'japan',
            LanguageCode.AR: 'ar',
        }
        return mapping.get(lang, 'en')


class LayoutAnalyzer:
    """Document layout analysis for structure preservation."""

    def __init__(self):
        """Initialize layout analyzer."""
        logger.info("Initialized LayoutAnalyzer")

    def analyze_layout(self, document_path: Path) -> LayoutAnalysis:
        """Analyze document layout to preserve structure.

        Args:
            document_path: Path to document file

        Returns:
            LayoutAnalysis with detected elements
        """
        try:
            logger.info(f"Analyzing layout for {document_path}")

            elements = []
            page_layouts = {}

            # Simple layout analysis implementation
            # In production, this would use more sophisticated layout detection
            if document_path.suffix.lower() == '.pdf':
                elements = self._analyze_pdf_layout(document_path)
            else:
                elements = self._analyze_image_layout(document_path)

            # Group elements by page
            for element in elements:
                page_num = element.page_number
                if page_num not in page_layouts:
                    page_layouts[page_num] = []
                page_layouts[page_num].append(element)

            # Generate reading order
            reading_order = self._determine_reading_order(elements)

            return LayoutAnalysis(
                document_path=str(document_path),
                elements=elements,
                page_layouts=page_layouts,
                reading_order=reading_order
            )

        except Exception as e:
            logger.error(f"Layout analysis failed for {document_path}: {e}")
            raise

    def _analyze_pdf_layout(self, document_path: Path) -> List[LayoutElement]:
        """Analyze PDF layout using PyMuPDF."""
        elements = []

        try:
            doc = fitz.open(str(document_path))

            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)

                # Get text blocks
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    if "lines" in block:  # Text block
                        bbox = block["bbox"]
                        text_content = ""

                        for line in block["lines"]:
                            for span in line["spans"]:
                                text_content += span["text"] + " "

                        if text_content.strip():
                            element = LayoutElement(
                                element_type="text_block",
                                bbox=(int(bbox[0]), int(bbox[1]), int(bbox[2] - bbox[0]), int(bbox[3] - bbox[1])),
                                page_number=page_num + 1,
                                content=text_content.strip(),
                                confidence=0.9
                            )
                            elements.append(element)

            doc.close()

        except Exception as e:
            logger.error(f"PDF layout analysis failed: {e}")

        return elements

    def _analyze_image_layout(self, document_path: Path) -> List[LayoutElement]:
        """Analyze image layout using computer vision."""
        elements = []

        try:
            # Load image
            image = cv2.imread(str(document_path))
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

            # Simple contour-based layout detection
            contours, _ = cv2.findContours(gray, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            for i, contour in enumerate(contours):
                x, y, w, h = cv2.boundingRect(contour)

                # Filter small contours
                if w > 50 and h > 20:
                    element = LayoutElement(
                        element_type="region",
                        bbox=(x, y, w, h),
                        page_number=1,
                        confidence=0.7
                    )
                    elements.append(element)

        except Exception as e:
            logger.error(f"Image layout analysis failed: {e}")

        return elements

    def _determine_reading_order(self, elements: List[LayoutElement]) -> List[str]:
        """Determine logical reading order of elements."""
        # Simple top-to-bottom, left-to-right ordering
        sorted_elements = sorted(elements, key=lambda e: (e.page_number, e.bbox[1], e.bbox[0]))
        return [f"element_{i}" for i in range(len(sorted_elements))]


class TextExtractor:
    """Text extraction from various document formats."""

    def __init__(self):
        """Initialize text extractor."""
        logger.info("Initialized TextExtractor")

    def extract_text(self, document_path: Path, doc_type: DocumentType) -> ExtractionResult:
        """Extract text from document using format-specific methods.

        Args:
            document_path: Path to document file
            doc_type: Type of document

        Returns:
            ExtractionResult with extracted content
        """
        start_time = datetime.utcnow()

        try:
            logger.info(f"Extracting text from {doc_type} document: {document_path}")

            extraction_methods = {
                DocumentType.PDF: self._extract_pdf_text,
                DocumentType.DOCX: self._extract_docx_text,
                DocumentType.XLSX: self._extract_xlsx_text,
                DocumentType.TXT: self._extract_txt_text,
            }

            if doc_type not in extraction_methods:
                raise ValueError(f"Text extraction not supported for {doc_type}")

            extracted_text, metadata, tables, images = extraction_methods[doc_type](document_path)

            processing_time = (datetime.utcnow() - start_time).total_seconds()

            return ExtractionResult(
                document_path=str(document_path),
                document_type=doc_type,
                extracted_text=extracted_text,
                metadata=metadata,
                tables=tables,
                images=images,
                extraction_method=f"native_{doc_type.value}",
                success=len(extracted_text.strip()) > 0,
                processing_time=processing_time
            )

        except Exception as e:
            logger.error(f"Text extraction failed for {document_path}: {e}")
            raise

    def _extract_pdf_text(self, document_path: Path) -> Tuple[str, DocumentMetadata, List[pd.DataFrame], List[Dict]]:
        """Extract text from PDF document."""
        doc = fitz.open(str(document_path))

        # Extract text
        text_content = []
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text_content.append(page.get_text())

        # Extract metadata
        metadata_dict = doc.metadata
        metadata = DocumentMetadata(
            title=metadata_dict.get('title'),
            author=metadata_dict.get('author'),
            subject=metadata_dict.get('subject'),
            creator=metadata_dict.get('creator'),
            producer=metadata_dict.get('producer'),
            page_count=doc.page_count
        )

        doc.close()

        return "\n".join(text_content), metadata, [], []

    def _extract_docx_text(self, document_path: Path) -> Tuple[str, DocumentMetadata, List[pd.DataFrame], List[Dict]]:
        """Extract text from DOCX document."""
        doc = Document(document_path)

        # Extract text from paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else [])
                tables.append(df)

        # Basic metadata
        metadata = DocumentMetadata(page_count=1)

        return "\n".join(text_content), metadata, tables, []

    def _extract_xlsx_text(self, document_path: Path) -> Tuple[str, DocumentMetadata, List[pd.DataFrame], List[Dict]]:
        """Extract text from XLSX document."""
        workbook = load_workbook(document_path, read_only=True)

        text_content = []
        tables = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            # Extract all cell values
            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data):  # Skip empty rows
                    sheet_data.append(row_data)
                    text_content.append(" ".join(row_data))

            # Create DataFrame for each sheet
            if sheet_data:
                df = pd.DataFrame(sheet_data)
                df.name = sheet_name
                tables.append(df)

        workbook.close()

        metadata = DocumentMetadata(page_count=len(workbook.sheetnames))
        return "\n".join(text_content), metadata, tables, []

    def _extract_txt_text(self, document_path: Path) -> Tuple[str, DocumentMetadata, List[pd.DataFrame], List[Dict]]:
        """Extract text from plain text file."""
        with open(document_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()

        metadata = DocumentMetadata(page_count=1)
        return content, metadata, [], []
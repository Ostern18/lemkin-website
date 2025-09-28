"""
Export Manager for Lemkin Report Generator Suite.

This module provides the ExportManager class for converting legal reports
to various formats (PDF, Word, LaTeX, HTML) with professional formatting,
quality validation, and court-ready output suitable for filing and distribution.
"""

import os
import io
import json
import hashlib
import tempfile
import subprocess
from datetime import datetime, date
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
import base64

from loguru import logger

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("ReportLab not available - PDF export will be limited")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word export will be limited")

from .core import (
    BaseReportModel, ReportConfig, ExportedReport, ExportSettings, 
    ExportFormat, FactSheet, EvidenceCatalog, LegalBrief, PersonInfo,
    CitationStyle, DocumentStandard
)


class PDFExporter:
    """Exports reports to PDF format using ReportLab"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.logger = logger.bind(component="pdf_exporter")
        
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
        
        # Initialize styles
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.black
        ))
        
        # Header style
        self.styles.add(ParagraphStyle(
            name='Header1',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.black,
            keepWithNext=True
        ))
        
        # Header style
        self.styles.add(ParagraphStyle(
            name='Header2',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceAfter=8,
            textColor=colors.black,
            keepWithNext=True
        ))
        
        # Body text with legal formatting
        self.styles.add(ParagraphStyle(
            name='LegalBody',
            parent=self.styles['Normal'],
            fontSize=12,
            leading=24,  # Double spacing
            firstLineIndent=0.5*inch,
            alignment=0  # Left justified
        ))
        
        # Citation style
        self.styles.add(ParagraphStyle(
            name='Citation',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=12,
            leftIndent=0.5*inch,
            rightIndent=0.5*inch
        ))
    
    def export_fact_sheet(
        self, 
        fact_sheet: FactSheet, 
        settings: ExportSettings
    ) -> Path:
        """Export fact sheet to PDF"""
        output_path = self._generate_output_path(settings, "fact_sheet", fact_sheet.case_data.case_info.case_number)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=settings.page_margins[0]*inch,
            leftMargin=settings.page_margins[1]*inch,
            topMargin=settings.page_margins[2]*inch,
            bottomMargin=settings.page_margins[3]*inch
        )
        
        story = []
        
        # Title
        story.append(Paragraph(fact_sheet.title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Case header
        case_info = fact_sheet.case_data.case_info
        header_data = [
            ['Case:', case_info.case_name],
            ['Case Number:', case_info.case_number],
            ['Court:', case_info.court],
            ['Judge:', case_info.judge or 'TBD'],
            ['Prepared By:', fact_sheet.author.full_name],
            ['Date:', fact_sheet.preparation_date.strftime('%B %d, %Y')]
        ]
        
        header_table = Table(header_data, colWidths=[1.5*inch, 4*inch])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(header_table)
        story.append(Spacer(1, 20))
        
        # Executive Summary
        if fact_sheet.executive_summary:
            story.append(Paragraph("EXECUTIVE SUMMARY", self.styles['Header1']))
            story.append(Paragraph(fact_sheet.executive_summary, self.styles['LegalBody']))
            story.append(Spacer(1, 16))
        
        # Factual Background
        if fact_sheet.factual_background:
            story.append(Paragraph("FACTUAL BACKGROUND", self.styles['Header1']))
            story.append(Paragraph(fact_sheet.factual_background, self.styles['LegalBody']))
            story.append(Spacer(1, 16))
        
        # Legal Issues
        if fact_sheet.legal_issues:
            story.append(Paragraph("LEGAL ISSUES", self.styles['Header1']))
            for i, issue in enumerate(fact_sheet.legal_issues, 1):
                story.append(Paragraph(f"{i}. {issue}", self.styles['LegalBody']))
            story.append(Spacer(1, 16))
        
        # Preliminary Analysis
        if fact_sheet.preliminary_analysis:
            story.append(Paragraph("PRELIMINARY ANALYSIS", self.styles['Header1']))
            story.append(Paragraph(fact_sheet.preliminary_analysis, self.styles['LegalBody']))
            story.append(Spacer(1, 16))
        
        # Recommendations
        if fact_sheet.recommendations:
            story.append(Paragraph("RECOMMENDATIONS", self.styles['Header1']))
            for i, rec in enumerate(fact_sheet.recommendations, 1):
                story.append(Paragraph(f"{i}. {rec}", self.styles['LegalBody']))
        
        # Build PDF
        doc.build(story)
        
        self.logger.info(f"Fact sheet exported to PDF: {output_path}")
        return output_path
    
    def export_evidence_catalog(
        self,
        catalog: EvidenceCatalog,
        settings: ExportSettings
    ) -> Path:
        """Export evidence catalog to PDF"""
        output_path = self._generate_output_path(settings, "evidence_catalog", catalog.case_data.case_info.case_number)
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=settings.page_margins[0]*inch,
            leftMargin=settings.page_margins[1]*inch,
            topMargin=settings.page_margins[2]*inch,
            bottomMargin=settings.page_margins[3]*inch
        )
        
        story = []
        
        # Title
        story.append(Paragraph(catalog.title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Catalog info
        catalog_data = [
            ['Case:', catalog.case_data.case_info.case_name],
            ['Case Number:', catalog.case_data.case_info.case_number],
            ['Custodian:', catalog.custodian.full_name],
            ['Catalog Date:', catalog.catalog_date.strftime('%B %d, %Y')],
            ['Total Items:', str(len(catalog.case_data.evidence_list))]
        ]
        
        catalog_table = Table(catalog_data, colWidths=[1.5*inch, 4*inch])
        catalog_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(catalog_table)
        story.append(Spacer(1, 20))
        
        # Evidence summary table
        if catalog.case_data.evidence_list:
            story.append(Paragraph("EVIDENCE INVENTORY", self.styles['Header1']))
            
            # Create evidence table
            evidence_data = [['ID', 'Title', 'Type', 'Date Collected', 'Custodian']]
            
            for evidence in catalog.case_data.evidence_list:
                row = [
                    evidence.evidence_id,
                    evidence.title[:30] + "..." if len(evidence.title) > 30 else evidence.title,
                    evidence.evidence_type.value,
                    evidence.date_collected.strftime('%m/%d/%Y') if evidence.date_collected else 'N/A',
                    evidence.custodian
                ]
                evidence_data.append(row)
            
            evidence_table = Table(evidence_data, colWidths=[0.8*inch, 2.2*inch, 1*inch, 1*inch, 1*inch])
            evidence_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(evidence_table)
        
        # Build PDF
        doc.build(story)
        
        self.logger.info(f"Evidence catalog exported to PDF: {output_path}")
        return output_path
    
    def export_legal_brief(
        self,
        brief: LegalBrief,
        settings: ExportSettings
    ) -> Path:
        """Export legal brief to PDF"""
        output_path = self._generate_output_path(settings, "legal_brief", brief.case_data.case_info.case_number)
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=settings.page_margins[0]*inch,
            leftMargin=settings.page_margins[1]*inch,
            topMargin=settings.page_margins[2]*inch,
            bottomMargin=settings.page_margins[3]*inch
        )
        
        story = []
        
        # Title page
        story.append(Paragraph(brief.case_data.case_info.court.upper(), self.styles['Header1']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(brief.case_data.case_info.case_name, self.styles['CustomTitle']))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Case No. {brief.case_data.case_info.case_number}", self.styles['Header2']))
        story.append(Spacer(1, 40))
        story.append(Paragraph(brief.title.upper(), self.styles['Header1']))
        story.append(Spacer(1, 60))
        
        # Attorney info
        attorney_info = [
            brief.author.full_name,
            brief.author.organization or "",
            brief.author.contact_email or "",
            brief.author.contact_phone or "",
            f"Attorney for [Party]"
        ]
        
        for line in attorney_info:
            if line:
                story.append(Paragraph(line, self.styles['Normal']))
        
        story.append(Spacer(1, 40))
        
        # Statement of Issues
        if brief.statement_of_issues:
            story.append(Paragraph("STATEMENT OF ISSUES", self.styles['Header1']))
            for i, issue in enumerate(brief.statement_of_issues, 1):
                story.append(Paragraph(f"{i}. {issue}", self.styles['LegalBody']))
            story.append(Spacer(1, 20))
        
        # Argument sections
        for section in brief.argument_sections:
            story.append(Paragraph(section.title.upper(), self.styles['Header1']))
            story.append(Paragraph(section.content, self.styles['LegalBody']))
            
            # Add subsections
            for subsection in section.subsections:
                story.append(Paragraph(subsection.title, self.styles['Header2']))
                story.append(Paragraph(subsection.content, self.styles['LegalBody']))
            
            story.append(Spacer(1, 16))
        
        # Conclusion
        if brief.conclusion:
            story.append(Paragraph("CONCLUSION", self.styles['Header1']))
            story.append(Paragraph(brief.conclusion, self.styles['LegalBody']))
            story.append(Spacer(1, 16))
        
        # Prayer for Relief
        if brief.prayer_for_relief:
            story.append(Paragraph(brief.prayer_for_relief, self.styles['LegalBody']))
        
        # Build PDF
        doc.build(story)
        
        self.logger.info(f"Legal brief exported to PDF: {output_path}")
        return output_path
    
    def _generate_output_path(
        self, 
        settings: ExportSettings, 
        report_type: str, 
        case_number: str
    ) -> Path:
        """Generate output file path"""
        if settings.output_directory:
            output_dir = Path(settings.output_directory)
        else:
            output_dir = Path.cwd()
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = settings.filename_template.format(
            report_type=report_type,
            case_number=case_number.replace(" ", "_"),
            date=timestamp
        )
        
        return output_dir / f"{filename}.pdf"


class WordExporter:
    """Exports reports to Microsoft Word format"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.logger = logger.bind(component="word_exporter")
        
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
    
    def export_fact_sheet(
        self,
        fact_sheet: FactSheet,
        settings: ExportSettings
    ) -> Path:
        """Export fact sheet to Word document"""
        output_path = self._generate_output_path(settings, "fact_sheet", fact_sheet.case_data.case_info.case_number)
        
        # Create document
        doc = Document()
        self._setup_document_styles(doc, settings)
        
        # Title
        title = doc.add_heading(fact_sheet.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Case header table
        case_info = fact_sheet.case_data.case_info
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        headers = ['Case:', 'Case Number:', 'Court:', 'Judge:', 'Prepared By:', 'Date:']
        values = [
            case_info.case_name,
            case_info.case_number,
            case_info.court,
            case_info.judge or 'TBD',
            fact_sheet.author.full_name,
            fact_sheet.preparation_date.strftime('%B %d, %Y')
        ]
        
        for i, (header, value) in enumerate(zip(headers, values)):
            table.cell(i, 0).text = header
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Executive Summary
        if fact_sheet.executive_summary:
            doc.add_heading('EXECUTIVE SUMMARY', level=1)
            doc.add_paragraph(fact_sheet.executive_summary)
        
        # Factual Background
        if fact_sheet.factual_background:
            doc.add_heading('FACTUAL BACKGROUND', level=1)
            doc.add_paragraph(fact_sheet.factual_background)
        
        # Legal Issues
        if fact_sheet.legal_issues:
            doc.add_heading('LEGAL ISSUES', level=1)
            for i, issue in enumerate(fact_sheet.legal_issues, 1):
                doc.add_paragraph(f"{i}. {issue}", style='List Number')
        
        # Preliminary Analysis
        if fact_sheet.preliminary_analysis:
            doc.add_heading('PRELIMINARY ANALYSIS', level=1)
            doc.add_paragraph(fact_sheet.preliminary_analysis)
        
        # Recommendations
        if fact_sheet.recommendations:
            doc.add_heading('RECOMMENDATIONS', level=1)
            for i, rec in enumerate(fact_sheet.recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
        
        # Save document
        doc.save(str(output_path))
        
        self.logger.info(f"Fact sheet exported to Word: {output_path}")
        return output_path
    
    def export_evidence_catalog(
        self,
        catalog: EvidenceCatalog,
        settings: ExportSettings
    ) -> Path:
        """Export evidence catalog to Word document"""
        output_path = self._generate_output_path(settings, "evidence_catalog", catalog.case_data.case_info.case_number)
        
        doc = Document()
        self._setup_document_styles(doc, settings)
        
        # Title
        title = doc.add_heading(catalog.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Catalog info table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        headers = ['Case:', 'Case Number:', 'Custodian:', 'Catalog Date:', 'Total Items:']
        values = [
            catalog.case_data.case_info.case_name,
            catalog.case_data.case_info.case_number,
            catalog.custodian.full_name,
            catalog.catalog_date.strftime('%B %d, %Y'),
            str(len(catalog.case_data.evidence_list))
        ]
        
        for i, (header, value) in enumerate(zip(headers, values)):
            table.cell(i, 0).text = header
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Evidence inventory
        if catalog.case_data.evidence_list:
            doc.add_heading('EVIDENCE INVENTORY', level=1)
            
            # Create evidence table
            evidence_table = doc.add_table(rows=1, cols=5)
            evidence_table.style = 'Table Grid'
            
            # Header row
            header_cells = evidence_table.rows[0].cells
            headers = ['Evidence ID', 'Title', 'Type', 'Date Collected', 'Custodian']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for evidence in catalog.case_data.evidence_list:
                row = evidence_table.add_row().cells
                row[0].text = evidence.evidence_id
                row[1].text = evidence.title[:50] + "..." if len(evidence.title) > 50 else evidence.title
                row[2].text = evidence.evidence_type.value
                row[3].text = evidence.date_collected.strftime('%m/%d/%Y') if evidence.date_collected else 'N/A'
                row[4].text = evidence.custodian
        
        # Save document
        doc.save(str(output_path))
        
        self.logger.info(f"Evidence catalog exported to Word: {output_path}")
        return output_path
    
    def export_legal_brief(
        self,
        brief: LegalBrief,
        settings: ExportSettings
    ) -> Path:
        """Export legal brief to Word document"""
        output_path = self._generate_output_path(settings, "legal_brief", brief.case_data.case_info.case_number)
        
        doc = Document()
        self._setup_document_styles(doc, settings)
        
        # Title page
        court_para = doc.add_paragraph(brief.case_data.case_info.court.upper())
        court_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        
        case_name = doc.add_heading(brief.case_data.case_info.case_name, level=1)
        case_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        case_num = doc.add_paragraph(f"Case No. {brief.case_data.case_info.case_number}")
        case_num.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        brief_title = doc.add_heading(brief.title.upper(), level=1)
        brief_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Statement of Issues
        if brief.statement_of_issues:
            doc.add_heading('STATEMENT OF ISSUES', level=1)
            for i, issue in enumerate(brief.statement_of_issues, 1):
                doc.add_paragraph(f"{i}. {issue}")
        
        # Argument sections
        for section in brief.argument_sections:
            doc.add_heading(section.title.upper(), level=1)
            doc.add_paragraph(section.content)
            
            # Add subsections
            for subsection in section.subsections:
                doc.add_heading(subsection.title, level=2)
                doc.add_paragraph(subsection.content)
        
        # Conclusion
        if brief.conclusion:
            doc.add_heading('CONCLUSION', level=1)
            doc.add_paragraph(brief.conclusion)
        
        # Prayer for Relief
        if brief.prayer_for_relief:
            doc.add_paragraph(brief.prayer_for_relief)
        
        # Save document
        doc.save(str(output_path))
        
        self.logger.info(f"Legal brief exported to Word: {output_path}")
        return output_path
    
    def _setup_document_styles(self, doc: Document, settings: ExportSettings):
        """Setup document styles and formatting"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = settings.font_family
        font.size = Pt(settings.font_size)
        
        # Set line spacing for legal documents
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = settings.line_spacing
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(settings.page_margins[2])
            section.bottom_margin = Inches(settings.page_margins[3])
            section.left_margin = Inches(settings.page_margins[1])
            section.right_margin = Inches(settings.page_margins[0])
    
    def _generate_output_path(
        self,
        settings: ExportSettings,
        report_type: str,
        case_number: str
    ) -> Path:
        """Generate output file path"""
        if settings.output_directory:
            output_dir = Path(settings.output_directory)
        else:
            output_dir = Path.cwd()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = settings.filename_template.format(
            report_type=report_type,
            case_number=case_number.replace(" ", "_"),
            date=timestamp
        )
        
        return output_dir / f"{filename}.docx"


class HTMLExporter:
    """Exports reports to HTML format"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.logger = logger.bind(component="html_exporter")
    
    def export_fact_sheet(
        self,
        fact_sheet: FactSheet,
        settings: ExportSettings
    ) -> Path:
        """Export fact sheet to HTML"""
        output_path = self._generate_output_path(settings, "fact_sheet", fact_sheet.case_data.case_info.case_number)
        
        html_content = self._generate_fact_sheet_html(fact_sheet, settings)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.logger.info(f"Fact sheet exported to HTML: {output_path}")
        return output_path
    
    def export_evidence_catalog(
        self,
        catalog: EvidenceCatalog,
        settings: ExportSettings
    ) -> Path:
        """Export evidence catalog to HTML"""
        output_path = self._generate_output_path(settings, "evidence_catalog", catalog.case_data.case_info.case_number)
        
        html_content = self._generate_catalog_html(catalog, settings)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.logger.info(f"Evidence catalog exported to HTML: {output_path}")
        return output_path
    
    def export_legal_brief(
        self,
        brief: LegalBrief,
        settings: ExportSettings
    ) -> Path:
        """Export legal brief to HTML"""
        output_path = self._generate_output_path(settings, "legal_brief", brief.case_data.case_info.case_number)
        
        html_content = self._generate_brief_html(brief, settings)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.logger.info(f"Legal brief exported to HTML: {output_path}")
        return output_path
    
    def _generate_fact_sheet_html(self, fact_sheet: FactSheet, settings: ExportSettings) -> str:
        """Generate HTML content for fact sheet"""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{fact_sheet.title}</title>
    <style>
        {self._get_base_css(settings)}
    </style>
</head>
<body>
    <div class="document">
        <h1 class="title">{fact_sheet.title}</h1>
        
        <table class="case-info">
            <tr><td class="label">Case:</td><td>{fact_sheet.case_data.case_info.case_name}</td></tr>
            <tr><td class="label">Case Number:</td><td>{fact_sheet.case_data.case_info.case_number}</td></tr>
            <tr><td class="label">Court:</td><td>{fact_sheet.case_data.case_info.court}</td></tr>
            <tr><td class="label">Judge:</td><td>{fact_sheet.case_data.case_info.judge or 'TBD'}</td></tr>
            <tr><td class="label">Prepared By:</td><td>{fact_sheet.author.full_name}</td></tr>
            <tr><td class="label">Date:</td><td>{fact_sheet.preparation_date.strftime('%B %d, %Y')}</td></tr>
        </table>
        
        """
        
        if fact_sheet.executive_summary:
            html += f"""
        <h2>EXECUTIVE SUMMARY</h2>
        <p class="content">{fact_sheet.executive_summary}</p>
        """
        
        if fact_sheet.factual_background:
            html += f"""
        <h2>FACTUAL BACKGROUND</h2>
        <p class="content">{fact_sheet.factual_background}</p>
        """
        
        if fact_sheet.legal_issues:
            html += "<h2>LEGAL ISSUES</h2><ol class='legal-list'>"
            for issue in fact_sheet.legal_issues:
                html += f"<li>{issue}</li>"
            html += "</ol>"
        
        if fact_sheet.preliminary_analysis:
            html += f"""
        <h2>PRELIMINARY ANALYSIS</h2>
        <p class="content">{fact_sheet.preliminary_analysis}</p>
        """
        
        if fact_sheet.recommendations:
            html += "<h2>RECOMMENDATIONS</h2><ol class='legal-list'>"
            for rec in fact_sheet.recommendations:
                html += f"<li>{rec}</li>"
            html += "</ol>"
        
        html += """
    </div>
</body>
</html>
        """
        
        return html
    
    def _generate_catalog_html(self, catalog: EvidenceCatalog, settings: ExportSettings) -> str:
        """Generate HTML content for evidence catalog"""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{catalog.title}</title>
    <style>
        {self._get_base_css(settings)}
        table.evidence {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
        table.evidence th, table.evidence td {{ border: 1px solid #333; padding: 8px; text-align: left; }}
        table.evidence th {{ background-color: #f0f0f0; font-weight: bold; }}
    </style>
</head>
<body>
    <div class="document">
        <h1 class="title">{catalog.title}</h1>
        
        <table class="case-info">
            <tr><td class="label">Case:</td><td>{catalog.case_data.case_info.case_name}</td></tr>
            <tr><td class="label">Case Number:</td><td>{catalog.case_data.case_info.case_number}</td></tr>
            <tr><td class="label">Custodian:</td><td>{catalog.custodian.full_name}</td></tr>
            <tr><td class="label">Catalog Date:</td><td>{catalog.catalog_date.strftime('%B %d, %Y')}</td></tr>
            <tr><td class="label">Total Items:</td><td>{len(catalog.case_data.evidence_list)}</td></tr>
        </table>
        
        <h2>EVIDENCE INVENTORY</h2>
        <table class="evidence">
            <tr>
                <th>Evidence ID</th>
                <th>Title</th>
                <th>Type</th>
                <th>Date Collected</th>
                <th>Custodian</th>
            </tr>
        """
        
        for evidence in catalog.case_data.evidence_list:
            html += f"""
            <tr>
                <td>{evidence.evidence_id}</td>
                <td>{evidence.title}</td>
                <td>{evidence.evidence_type.value}</td>
                <td>{evidence.date_collected.strftime('%m/%d/%Y') if evidence.date_collected else 'N/A'}</td>
                <td>{evidence.custodian}</td>
            </tr>
            """
        
        html += """
        </table>
    </div>
</body>
</html>
        """
        
        return html
    
    def _generate_brief_html(self, brief: LegalBrief, settings: ExportSettings) -> str:
        """Generate HTML content for legal brief"""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{brief.title}</title>
    <style>
        {self._get_base_css(settings)}
    </style>
</head>
<body>
    <div class="document">
        <h1 class="court">{brief.case_data.case_info.court.upper()}</h1>
        <h1 class="case-name">{brief.case_data.case_info.case_name}</h1>
        <h2 class="case-number">Case No. {brief.case_data.case_info.case_number}</h2>
        <h1 class="brief-title">{brief.title.upper()}</h1>
        """
        
        if brief.statement_of_issues:
            html += "<h2>STATEMENT OF ISSUES</h2><ol class='legal-list'>"
            for issue in brief.statement_of_issues:
                html += f"<li>{issue}</li>"
            html += "</ol>"
        
        # Argument sections
        for section in brief.argument_sections:
            html += f"<h2>{section.title.upper()}</h2>"
            html += f"<p class='content'>{section.content}</p>"
            
            for subsection in section.subsections:
                html += f"<h3>{subsection.title}</h3>"
                html += f"<p class='content'>{subsection.content}</p>"
        
        if brief.conclusion:
            html += f"""
        <h2>CONCLUSION</h2>
        <p class="content">{brief.conclusion}</p>
        """
        
        if brief.prayer_for_relief:
            html += f"""
        <p class="prayer">{brief.prayer_for_relief}</p>
        """
        
        html += """
    </div>
</body>
</html>
        """
        
        return html
    
    def _get_base_css(self, settings: ExportSettings) -> str:
        """Generate base CSS for HTML documents"""
        return f"""
        body {{
            font-family: {settings.font_family}, serif;
            font-size: {settings.font_size}pt;
            line-height: {settings.line_spacing};
            margin: {settings.page_margins[2]}in {settings.page_margins[0]}in {settings.page_margins[3]}in {settings.page_margins[1]}in;
            color: #000;
        }}
        .document {{
            max-width: 8.5in;
            margin: 0 auto;
        }}
        .title, .case-name, .brief-title {{
            text-align: center;
            margin: 20px 0;
        }}
        .court, .case-number {{
            text-align: center;
            margin: 10px 0;
        }}
        h1 {{
            font-size: 16pt;
            font-weight: bold;
            margin: 20px 0 10px 0;
        }}
        h2 {{
            font-size: 14pt;
            font-weight: bold;
            margin: 16px 0 8px 0;
        }}
        h3 {{
            font-size: 12pt;
            font-weight: bold;
            margin: 12px 0 6px 0;
        }}
        .content {{
            text-align: justify;
            margin: 10px 0;
            text-indent: 0.5in;
        }}
        table.case-info {{
            margin: 20px 0;
            border-collapse: collapse;
        }}
        table.case-info td {{
            padding: 5px 10px;
            border: none;
        }}
        .label {{
            font-weight: bold;
            width: 150px;
        }}
        .legal-list {{
            margin: 10px 0;
            padding-left: 20px;
        }}
        .legal-list li {{
            margin: 8px 0;
        }}
        .prayer {{
            margin: 20px 0;
            white-space: pre-line;
        }}
        """
    
    def _generate_output_path(
        self,
        settings: ExportSettings,
        report_type: str,
        case_number: str
    ) -> Path:
        """Generate output file path"""
        if settings.output_directory:
            output_dir = Path(settings.output_directory)
        else:
            output_dir = Path.cwd()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = settings.filename_template.format(
            report_type=report_type,
            case_number=case_number.replace(" ", "_"),
            date=timestamp
        )
        
        return output_dir / f"{filename}.html"


class LaTeXExporter:
    """Exports reports to LaTeX format for professional typesetting"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.logger = logger.bind(component="latex_exporter")
    
    def export_legal_brief(
        self,
        brief: LegalBrief,
        settings: ExportSettings
    ) -> Path:
        """Export legal brief to LaTeX format"""
        output_path = self._generate_output_path(settings, "legal_brief", brief.case_data.case_info.case_number)
        
        latex_content = self._generate_brief_latex(brief, settings)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        self.logger.info(f"Legal brief exported to LaTeX: {output_path}")
        return output_path
    
    def _generate_brief_latex(self, brief: LegalBrief, settings: ExportSettings) -> str:
        """Generate LaTeX content for legal brief"""
        latex = r"""
\documentclass[12pt]{article}
\usepackage[letterpaper,margin=1in]{geometry}
\usepackage{setspace}
\doublespacing
\usepackage{indentfirst}
\setlength{\parindent}{0.5in}

\begin{document}

"""
        
        # Title page
        latex += f"""
\\begin{{center}}
{{\\large\\textbf{{{brief.case_data.case_info.court.upper()}}}}}

\\vspace{{1cm}}

{{\\Large\\textbf{{{brief.case_data.case_info.case_name}}}}}

\\vspace{{0.5cm}}

Case No. {brief.case_data.case_info.case_number}

\\vspace{{2cm}}

{{\\large\\textbf{{{brief.title.upper()}}}}}

\\vspace{{3cm}}

{brief.author.full_name} \\\\
{brief.author.organization or ""} \\\\
{brief.author.contact_email or ""} \\\\
Attorney for [Party]

\\end{{center}}

\\newpage

"""
        
        # Statement of Issues
        if brief.statement_of_issues:
            latex += r"\section*{STATEMENT OF ISSUES}" + "\n\n"
            latex += r"\begin{enumerate}" + "\n"
            for issue in brief.statement_of_issues:
                latex += f"\\item {self._escape_latex(issue)}\n\n"
            latex += r"\end{enumerate}" + "\n\n"
        
        # Argument sections
        for section in brief.argument_sections:
            latex += f"\\section*{{{section.title.upper()}}}\n\n"
            latex += f"{self._escape_latex(section.content)}\n\n"
            
            for subsection in section.subsections:
                latex += f"\\subsection*{{{subsection.title}}}\n\n"
                latex += f"{self._escape_latex(subsection.content)}\n\n"
        
        # Conclusion
        if brief.conclusion:
            latex += r"\section*{CONCLUSION}" + "\n\n"
            latex += f"{self._escape_latex(brief.conclusion)}\n\n"
        
        # Prayer for Relief
        if brief.prayer_for_relief:
            latex += f"{self._escape_latex(brief.prayer_for_relief)}\n\n"
        
        latex += r"\end{document}"
        
        return latex
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        replacements = [
            ('\\', r'\textbackslash{}'),
            ('{', r'\{'),
            ('}', r'\}'),
            ('$', r'\$'),
            ('&', r'\&'),
            ('%', r'\%'),
            ('#', r'\#'),
            ('^', r'\textasciicircum{}'),
            ('_', r'\_'),
            ('~', r'\textasciitilde{}'),
        ]
        
        for old, new in replacements:
            text = text.replace(old, new)
        
        return text
    
    def _generate_output_path(
        self,
        settings: ExportSettings,
        report_type: str,
        case_number: str
    ) -> Path:
        """Generate output file path"""
        if settings.output_directory:
            output_dir = Path(settings.output_directory)
        else:
            output_dir = Path.cwd()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = settings.filename_template.format(
            report_type=report_type,
            case_number=case_number.replace(" ", "_"),
            date=timestamp
        )
        
        return output_dir / f"{filename}.tex"


class ExportManager:
    """
    Multi-format report export manager.
    
    Provides professional document export capabilities for legal reports
    with quality validation, format-specific optimization, and court-ready
    output suitable for filing and distribution.
    """
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.logger = logger.bind(component="export_manager")
        
        # Initialize exporters
        self._exporters = {}
        
        try:
            self._exporters[ExportFormat.PDF] = PDFExporter(config)
        except ImportError:
            self.logger.warning("PDF export not available - install reportlab")
        
        try:
            self._exporters[ExportFormat.DOCX] = WordExporter(config)
        except ImportError:
            self.logger.warning("Word export not available - install python-docx")
        
        self._exporters[ExportFormat.HTML] = HTMLExporter(config)
        self._exporters[ExportFormat.LATEX] = LaTeXExporter(config)
    
    def export(
        self,
        report: Union[FactSheet, EvidenceCatalog, LegalBrief],
        format_type: ExportFormat,
        settings: Optional[ExportSettings] = None,
        exported_by: Optional[PersonInfo] = None
    ) -> ExportedReport:
        """
        Export report to specified format with quality validation
        
        Args:
            report: Report to export
            format_type: Target export format
            settings: Custom export settings
            exported_by: Person exporting the report
            
        Returns:
            ExportedReport with file information and metadata
        """
        start_time = datetime.utcnow()
        
        # Use default settings if not provided
        if not settings:
            settings = ExportSettings(export_format=format_type)
        
        # Get default exporter if not provided
        if not exported_by:
            exported_by = PersonInfo(
                full_name="System Export",
                role="system",
                organization=self.config.firm_name
            )
        
        self.logger.info(f"Exporting {type(report).__name__} to {format_type.value}")
        
        try:
            # Check if exporter is available
            if format_type not in self._exporters:
                raise ValueError(f"Export format {format_type.value} not available")
            
            exporter = self._exporters[format_type]
            
            # Export based on report type
            if isinstance(report, FactSheet):
                output_path = exporter.export_fact_sheet(report, settings)
            elif isinstance(report, EvidenceCatalog):
                output_path = exporter.export_evidence_catalog(report, settings)
            elif isinstance(report, LegalBrief):
                output_path = exporter.export_legal_brief(report, settings)
            else:
                raise ValueError(f"Unsupported report type: {type(report)}")
            
            # Calculate file hash
            file_size = output_path.stat().st_size
            with open(output_path, 'rb') as f:
                file_content = f.read()
                file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Validate export
            validation_result = self._validate_export(output_path, format_type)
            
            # Create export record
            exported_report = ExportedReport(
                source_report=report,
                export_settings=settings,
                exported_by=exported_by,
                export_format=format_type,
                output_path=output_path,
                file_size_bytes=file_size,
                file_hash=file_hash,
                export_successful=True,
                validation_passed=validation_result["valid"],
                compliance_check_passed=validation_result.get("compliant", True)
            )
            
            if validation_result["errors"]:
                exported_report.export_errors.extend(validation_result["errors"])
            
            if validation_result["warnings"]:
                exported_report.warnings.extend(validation_result["warnings"])
            
            duration = (datetime.utcnow() - start_time).total_seconds()
            self.logger.info(f"Export completed successfully in {duration:.2f} seconds")
            
        except Exception as e:
            self.logger.error(f"Export failed: {str(e)}")
            
            # Create failed export record
            exported_report = ExportedReport(
                source_report=report,
                export_settings=settings,
                exported_by=exported_by,
                export_format=format_type,
                output_path=Path("/dev/null"),  # Placeholder
                file_size_bytes=0,
                file_hash="",
                export_successful=False,
                validation_passed=False,
                compliance_check_passed=False
            )
            
            exported_report.export_errors.append(str(e))
            raise
        
        return exported_report
    
    def batch_export(
        self,
        reports: List[Union[FactSheet, EvidenceCatalog, LegalBrief]],
        formats: List[ExportFormat],
        settings: Optional[ExportSettings] = None
    ) -> List[ExportedReport]:
        """
        Export multiple reports to multiple formats
        
        Args:
            reports: List of reports to export
            formats: List of target formats
            settings: Export settings to use for all exports
            
        Returns:
            List of ExportedReport objects
        """
        exported_reports = []
        
        self.logger.info(f"Starting batch export: {len(reports)} reports to {len(formats)} formats")
        
        for report in reports:
            for format_type in formats:
                try:
                    exported_report = self.export(report, format_type, settings)
                    exported_reports.append(exported_report)
                except Exception as e:
                    self.logger.error(f"Batch export failed for {type(report).__name__} to {format_type.value}: {str(e)}")
                    continue
        
        self.logger.info(f"Batch export completed: {len(exported_reports)} files generated")
        return exported_reports
    
    def _validate_export(self, output_path: Path, format_type: ExportFormat) -> Dict[str, Any]:
        """Validate exported file"""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "compliant": True
        }
        
        # Check file exists and has content
        if not output_path.exists():
            validation_result["valid"] = False
            validation_result["errors"].append("Export file does not exist")
            return validation_result
        
        if output_path.stat().st_size == 0:
            validation_result["valid"] = False
            validation_result["errors"].append("Export file is empty")
            return validation_result
        
        # Format-specific validation
        if format_type == ExportFormat.PDF:
            validation_result.update(self._validate_pdf(output_path))
        elif format_type == ExportFormat.DOCX:
            validation_result.update(self._validate_docx(output_path))
        elif format_type == ExportFormat.HTML:
            validation_result.update(self._validate_html(output_path))
        
        return validation_result
    
    def _validate_pdf(self, pdf_path: Path) -> Dict[str, Any]:
        """Validate PDF file"""
        validation = {"valid": True, "errors": [], "warnings": []}
        
        try:
            # Basic validation - check if file starts with PDF header
            with open(pdf_path, 'rb') as f:
                header = f.read(4)
                if not header.startswith(b'%PDF'):
                    validation["valid"] = False
                    validation["errors"].append("Invalid PDF file header")
        except Exception as e:
            validation["valid"] = False
            validation["errors"].append(f"PDF validation error: {str(e)}")
        
        return validation
    
    def _validate_docx(self, docx_path: Path) -> Dict[str, Any]:
        """Validate Word document"""
        validation = {"valid": True, "errors": [], "warnings": []}
        
        try:
            # Check if it's a valid ZIP file (DOCX format)
            import zipfile
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # Check for required DOCX structure
                required_files = ['word/document.xml', '[Content_Types].xml']
                for required_file in required_files:
                    if required_file not in zip_file.namelist():
                        validation["warnings"].append(f"Missing {required_file} in DOCX structure")
        except Exception as e:
            validation["valid"] = False
            validation["errors"].append(f"DOCX validation error: {str(e)}")
        
        return validation
    
    def _validate_html(self, html_path: Path) -> Dict[str, Any]:
        """Validate HTML file"""
        validation = {"valid": True, "errors": [], "warnings": []}
        
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
                # Basic HTML validation
                if not content.strip().startswith('<!DOCTYPE html>'):
                    validation["warnings"].append("Missing DOCTYPE declaration")
                
                if '<html>' not in content:
                    validation["errors"].append("Missing HTML element")
                    validation["valid"] = False
                
                if '<title>' not in content:
                    validation["warnings"].append("Missing title element")
        
        except Exception as e:
            validation["valid"] = False
            validation["errors"].append(f"HTML validation error: {str(e)}")
        
        return validation
    
    def get_supported_formats(self) -> List[ExportFormat]:
        """Get list of supported export formats"""
        return list(self._exporters.keys())
    
    def is_format_supported(self, format_type: ExportFormat) -> bool:
        """Check if export format is supported"""
        return format_type in self._exporters
    
    def get_export_statistics(self) -> Dict[str, Any]:
        """Get export statistics and capabilities"""
        return {
            "supported_formats": [fmt.value for fmt in self._exporters.keys()],
            "pdf_available": ExportFormat.PDF in self._exporters,
            "word_available": ExportFormat.DOCX in self._exporters,
            "html_available": ExportFormat.HTML in self._exporters,
            "latex_available": ExportFormat.LATEX in self._exporters
        }
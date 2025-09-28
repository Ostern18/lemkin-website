"""
Core functionality for legal document classification using fine-tuned BERT models.

This module provides the main DocumentClassifier class and supporting data models
for automated classification of legal documents in various formats.
"""

import json
import pickle
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple
from dataclasses import asdict

import numpy as np
import pandas as pd
import torch
from transformers import (
    AutoTokenizer, 
    AutoModelForSequenceClassification,
    Trainer,
    TrainingArguments,
    pipeline,
)
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score
from sklearn.model_selection import train_test_split
from datasets import Dataset
from pydantic import BaseModel, Field, validator
import pdfplumber
from docx import Document as DocxDocument

from .legal_taxonomy import (
    DocumentType, 
    LegalDomain, 
    LegalDocumentCategory,
    get_category_hierarchy,
    validate_category,
    get_supported_categories,
)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentContent(BaseModel):
    """Structured representation of document content for classification"""
    
    text: str = Field(description="Primary text content")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Document metadata")
    file_path: Optional[str] = Field(default=None, description="Original file path")
    file_type: str = Field(default="text", description="File type: text, pdf, docx, html")
    language: str = Field(default="en", description="Document language")
    length: int = Field(description="Text length in characters")
    word_count: int = Field(description="Word count")
    
    @validator('length', always=True)
    def set_length(cls, v, values):
        if 'text' in values:
            return len(values['text'])
        return 0
    
    @validator('word_count', always=True)
    def set_word_count(cls, v, values):
        if 'text' in values:
            return len(values['text'].split())
        return 0


class DocumentClassification(BaseModel):
    """Classification result for a document"""
    
    document_type: DocumentType = Field(description="Predicted document type")
    legal_domain: LegalDomain = Field(description="Predicted legal domain")
    confidence_score: float = Field(description="Classification confidence (0-1)")
    probability_distribution: Dict[str, float] = Field(description="Probability for each class")
    subcategory: Optional[str] = Field(default=None, description="Specific subcategory if applicable")
    
    # Model information
    model_version: str = Field(description="Version of the classification model used")
    model_name: str = Field(description="Name of the model used")
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    
    class Config:
        use_enum_values = True
        json_encoders = {
            datetime: lambda v: v.isoformat()
        }


class ClassificationResult(BaseModel):
    """Complete classification result with metadata and recommendations"""
    
    document_content: DocumentContent = Field(description="Original document content")
    classification: DocumentClassification = Field(description="Classification results")
    legal_category: LegalDocumentCategory = Field(description="Detailed category information")
    
    # Processing metadata
    processing_time: float = Field(description="Processing time in seconds")
    requires_review: bool = Field(description="Whether human review is recommended")
    review_reasons: List[str] = Field(default_factory=list, description="Reasons for review")
    
    # Recommendations
    recommended_actions: List[str] = Field(default_factory=list, description="Recommended next actions")
    urgency_level: str = Field(description="Processing urgency level")
    sensitivity_level: str = Field(description="Document sensitivity level")
    
    class Config:
        use_enum_values = True
        json_encoders = {
            datetime: lambda v: v.isoformat()
        }


class ClassificationConfig(BaseModel):
    """Configuration for document classification"""
    
    model_name: str = Field(default="distilbert-base-uncased", description="Base model name")
    model_path: Optional[str] = Field(default=None, description="Path to fine-tuned model")
    max_length: int = Field(default=512, description="Maximum sequence length")
    confidence_threshold: float = Field(default=0.7, description="Minimum confidence threshold")
    batch_size: int = Field(default=16, description="Batch size for processing")
    device: str = Field(default="auto", description="Computing device: auto, cpu, cuda")
    
    # Language settings
    supported_languages: List[str] = Field(default=["en"], description="Supported languages")
    default_language: str = Field(default="en", description="Default language")
    
    # Processing options
    enable_preprocessing: bool = Field(default=True, description="Enable text preprocessing")
    enable_multilingual: bool = Field(default=False, description="Enable multilingual support")
    cache_predictions: bool = Field(default=True, description="Cache prediction results")
    
    @validator('device')
    def validate_device(cls, v):
        if v == "auto":
            return "cuda" if torch.cuda.is_available() else "cpu"
        return v


class ModelMetrics(BaseModel):
    """Model performance metrics"""
    
    accuracy: float = Field(description="Overall accuracy")
    precision: float = Field(description="Macro-averaged precision")
    recall: float = Field(description="Macro-averaged recall")
    f1_score: float = Field(description="Macro-averaged F1 score")
    
    # Per-class metrics
    class_metrics: Dict[str, Dict[str, float]] = Field(description="Per-class precision/recall/f1")
    confusion_matrix: List[List[int]] = Field(description="Confusion matrix")
    
    # Model information
    model_version: str = Field(description="Model version")
    evaluation_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    test_set_size: int = Field(description="Size of test set")
    
    class Config:
        json_encoders = {
            datetime: lambda v: v.isoformat(),
            np.ndarray: lambda v: v.tolist(),
        }


class TrainingMetrics(BaseModel):
    """Training progress and metrics"""
    
    epoch: int = Field(description="Current epoch")
    train_loss: float = Field(description="Training loss")
    eval_loss: float = Field(description="Evaluation loss")
    eval_accuracy: float = Field(description="Evaluation accuracy")
    learning_rate: float = Field(description="Current learning rate")
    
    # Training metadata
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    total_steps: int = Field(description="Total training steps")
    
    class Config:
        json_encoders = {
            datetime: lambda v: v.isoformat()
        }


class DocumentClassifier:
    """Main class for legal document classification using BERT models"""
    
    def __init__(self, config: ClassificationConfig):
        """
        Initialize the document classifier
        
        Args:
            config: Classification configuration
        """
        self.config = config
        self.tokenizer = None
        self.model = None
        self.classifier_pipeline = None
        self.label_encoder = None
        self.class_labels = []
        
        # Initialize category hierarchy
        self.category_hierarchy = get_category_hierarchy()
        self.supported_categories = get_supported_categories()
        
        # Load or initialize model
        self._initialize_model()
        
        logger.info(f"DocumentClassifier initialized with model: {config.model_name}")
    
    def _initialize_model(self) -> None:
        """Initialize BERT model and tokenizer"""
        try:
            # Load tokenizer
            model_path = self.config.model_path or self.config.model_name
            self.tokenizer = AutoTokenizer.from_pretrained(model_path)
            
            # Load model
            if self.config.model_path and Path(self.config.model_path).exists():
                # Load fine-tuned model
                self.model = AutoModelForSequenceClassification.from_pretrained(
                    self.config.model_path
                )
                # Load class labels
                label_file = Path(self.config.model_path) / "class_labels.json"
                if label_file.exists():
                    with open(label_file, 'r') as f:
                        self.class_labels = json.load(f)
            else:
                # Initialize base model for training
                num_labels = len(self.supported_categories)
                self.model = AutoModelForSequenceClassification.from_pretrained(
                    self.config.model_name,
                    num_labels=num_labels
                )
                self.class_labels = [cat.value for cat in self.supported_categories]
            
            # Create classification pipeline
            self.classifier_pipeline = pipeline(
                "text-classification",
                model=self.model,
                tokenizer=self.tokenizer,
                device=0 if self.config.device == "cuda" else -1,
                return_all_scores=True
            )
            
            logger.info("Model and tokenizer loaded successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize model: {e}")
            raise
    
    def extract_text_from_file(self, file_path: Path) -> DocumentContent:
        """
        Extract text content from various file formats
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentContent object with extracted text and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text = ""
        file_type = file_path.suffix.lower()
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "created_date": datetime.fromtimestamp(file_path.stat().st_ctime),
            "modified_date": datetime.fromtimestamp(file_path.stat().st_mtime),
        }
        
        try:
            if file_type == ".pdf":
                # Extract from PDF
                with pdfplumber.open(file_path) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                    metadata["num_pages"] = len(pdf.pages)
                file_type = "pdf"
                
            elif file_type == ".docx":
                # Extract from Word document
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                metadata["num_paragraphs"] = len(doc.paragraphs)
                file_type = "docx"
                
            elif file_type in [".txt", ".md"]:
                # Extract from text files
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                file_type = "text"
                
            else:
                # Try to read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                file_type = "text"
        
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise
        
        # Clean and validate text
        text = text.strip()
        if not text:
            raise ValueError(f"No text content extracted from {file_path}")
        
        return DocumentContent(
            text=text,
            metadata=metadata,
            file_path=str(file_path),
            file_type=file_type,
            language=self.config.default_language,
            length=len(text),
            word_count=len(text.split())
        )
    
    def preprocess_text(self, text: str) -> str:
        """
        Preprocess text for classification
        
        Args:
            text: Raw text content
            
        Returns:
            Preprocessed text
        """
        if not self.config.enable_preprocessing:
            return text
        
        # Basic text cleaning
        text = text.strip()
        
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Truncate if too long
        if len(text) > self.config.max_length * 4:  # Rough character limit
            text = text[:self.config.max_length * 4]
            logger.warning(f"Text truncated to {len(text)} characters")
        
        return text
    
    def classify_document(self, document_content: DocumentContent) -> ClassificationResult:
        """
        Classify a document and return detailed results
        
        Args:
            document_content: Document content to classify
            
        Returns:
            ClassificationResult with predictions and recommendations
        """
        start_time = datetime.now()
        
        # Preprocess text
        processed_text = self.preprocess_text(document_content.text)
        
        try:
            # Get predictions from model
            predictions = self.classifier_pipeline(processed_text)
            
            # Parse results
            prediction_scores = {pred['label']: pred['score'] for pred in predictions}
            
            # Get best prediction
            best_prediction = max(predictions, key=lambda x: x['score'])
            predicted_type = DocumentType(best_prediction['label'])
            confidence = best_prediction['score']
            
            # Get legal domain (simplified - could be enhanced with separate domain classifier)
            legal_domain = self._infer_legal_domain(predicted_type, document_content)
            
            # Create classification result
            classification = DocumentClassification(
                document_type=predicted_type,
                legal_domain=legal_domain,
                confidence_score=confidence,
                probability_distribution=prediction_scores,
                model_version="1.0",
                model_name=self.config.model_name
            )
            
            # Get category details
            legal_category = self.category_hierarchy.primary_categories.get(
                predicted_type,
                LegalDocumentCategory(
                    document_type=predicted_type,
                    legal_domain=legal_domain
                )
            )
            
            # Determine review requirements
            requires_review, review_reasons = self._assess_review_requirements(
                classification, legal_category, document_content
            )
            
            # Generate recommendations
            recommendations = self._generate_recommendations(
                classification, legal_category, document_content
            )
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            result = ClassificationResult(
                document_content=document_content,
                classification=classification,
                legal_category=legal_category,
                processing_time=processing_time,
                requires_review=requires_review,
                review_reasons=review_reasons,
                recommended_actions=recommendations,
                urgency_level=legal_category.urgency_level,
                sensitivity_level=legal_category.sensitivity_level
            )
            
            logger.info(
                f"Document classified as {predicted_type.value} "
                f"with confidence {confidence:.3f}"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Classification failed: {e}")
            raise
    
    def classify_file(self, file_path: Union[str, Path]) -> ClassificationResult:
        """
        Classify a document file
        
        Args:
            file_path: Path to document file
            
        Returns:
            ClassificationResult with predictions and recommendations
        """
        # Extract text content
        document_content = self.extract_text_from_file(Path(file_path))
        
        # Classify document
        return self.classify_document(document_content)
    
    def _infer_legal_domain(self, document_type: DocumentType, 
                           document_content: DocumentContent) -> LegalDomain:
        """
        Infer legal domain based on document type and content
        
        Args:
            document_type: Predicted document type
            document_content: Document content
            
        Returns:
            Inferred legal domain
        """
        # Simple domain inference based on document type
        # In practice, this could use a separate domain classifier
        
        domain_mapping = {
            DocumentType.POLICE_REPORT: LegalDomain.CRIMINAL_LAW,
            DocumentType.WITNESS_STATEMENT: LegalDomain.CRIMINAL_LAW,
            DocumentType.FORENSIC_REPORT: LegalDomain.CRIMINAL_LAW,
            DocumentType.MEDICAL_RECORD: LegalDomain.CIVIL_RIGHTS,
            DocumentType.MILITARY_REPORT: LegalDomain.INTERNATIONAL_HUMANITARIAN_LAW,
            DocumentType.GOVERNMENT_DOCUMENT: LegalDomain.ADMINISTRATIVE_LAW,
            DocumentType.FINANCIAL_RECORD: LegalDomain.CORPORATE_LAW,
            DocumentType.CONTRACT: LegalDomain.CORPORATE_LAW,
        }
        
        return domain_mapping.get(document_type, LegalDomain.GENERAL)
    
    def _assess_review_requirements(self, classification: DocumentClassification,
                                   legal_category: LegalDocumentCategory,
                                   document_content: DocumentContent) -> Tuple[bool, List[str]]:
        """
        Assess whether human review is required
        
        Args:
            classification: Classification results
            legal_category: Legal category information
            document_content: Original document content
            
        Returns:
            Tuple of (requires_review, review_reasons)
        """
        requires_review = False
        reasons = []
        
        # Check confidence threshold
        if classification.confidence_score < self.config.confidence_threshold:
            requires_review = True
            reasons.append(f"Low confidence score: {classification.confidence_score:.3f}")
        
        # Check category requirements
        if legal_category.requires_human_review:
            requires_review = True
            reasons.append(f"Document type {classification.document_type.value} requires review")
        
        # Check sensitivity level
        if legal_category.sensitivity_level in ["confidential", "restricted"]:
            requires_review = True
            reasons.append(f"High sensitivity level: {legal_category.sensitivity_level}")
        
        # Check urgency level
        if legal_category.urgency_level == "critical":
            requires_review = True
            reasons.append("Critical urgency level")
        
        # Check document characteristics
        if document_content.word_count < 50:
            requires_review = True
            reasons.append("Very short document content")
        
        return requires_review, reasons
    
    def _generate_recommendations(self, classification: DocumentClassification,
                                 legal_category: LegalDocumentCategory,
                                 document_content: DocumentContent) -> List[str]:
        """
        Generate processing recommendations
        
        Args:
            classification: Classification results
            legal_category: Legal category information
            document_content: Original document content
            
        Returns:
            List of recommended actions
        """
        recommendations = []
        
        # Redaction recommendations
        if legal_category.redaction_required:
            recommendations.append("Apply PII redaction before sharing")
        
        # Chain of custody recommendations
        if legal_category.chain_of_custody_critical:
            recommendations.append("Maintain strict chain of custody")
        
        # Urgency-based recommendations
        if legal_category.urgency_level == "critical":
            recommendations.append("Prioritize immediate processing")
        elif legal_category.urgency_level == "high":
            recommendations.append("Process within 24 hours")
        
        # Document-specific recommendations
        if classification.document_type == DocumentType.WITNESS_STATEMENT:
            recommendations.append("Verify witness identity and statement authenticity")
        elif classification.document_type == DocumentType.FORENSIC_REPORT:
            recommendations.append("Validate forensic methodology and chain of custody")
        elif classification.document_type == DocumentType.MEDICAL_RECORD:
            recommendations.append("Ensure HIPAA compliance for medical information")
        
        # Quality recommendations
        if classification.confidence_score < 0.8:
            recommendations.append("Consider manual verification of classification")
        
        return recommendations
    
    def train_model(self, training_data: List[Tuple[str, str]], 
                   validation_split: float = 0.2,
                   output_dir: Optional[str] = None) -> ModelMetrics:
        """
        Fine-tune the model on legal document training data
        
        Args:
            training_data: List of (text, label) tuples
            validation_split: Fraction of data for validation
            output_dir: Directory to save the trained model
            
        Returns:
            ModelMetrics with training results
        """
        if not training_data:
            raise ValueError("Training data cannot be empty")
        
        # Prepare data
        texts, labels = zip(*training_data)
        unique_labels = list(set(labels))
        
        # Create label mapping
        label_to_id = {label: i for i, label in enumerate(unique_labels)}
        id_to_label = {i: label for label, i in label_to_id.items()}
        
        # Convert labels to ids
        label_ids = [label_to_id[label] for label in labels]
        
        # Split data
        train_texts, val_texts, train_labels, val_labels = train_test_split(
            texts, label_ids, test_size=validation_split, random_state=42, stratify=label_ids
        )
        
        # Create datasets
        train_dataset = Dataset.from_dict({
            'text': train_texts,
            'labels': train_labels
        })
        
        val_dataset = Dataset.from_dict({
            'text': val_texts,
            'labels': val_labels
        })
        
        # Tokenize datasets
        def tokenize_function(examples):
            return self.tokenizer(
                examples['text'],
                truncation=True,
                padding=True,
                max_length=self.config.max_length
            )
        
        train_dataset = train_dataset.map(tokenize_function, batched=True)
        val_dataset = val_dataset.map(tokenize_function, batched=True)
        
        # Update model for new labels if needed
        if len(unique_labels) != self.model.config.num_labels:
            self.model = AutoModelForSequenceClassification.from_pretrained(
                self.config.model_name,
                num_labels=len(unique_labels)
            )
        
        # Training arguments
        training_args = TrainingArguments(
            output_dir=output_dir or "./results",
            num_train_epochs=3,
            per_device_train_batch_size=self.config.batch_size,
            per_device_eval_batch_size=self.config.batch_size,
            warmup_steps=500,
            weight_decay=0.01,
            logging_dir=f"{output_dir or './results'}/logs",
            evaluation_strategy="epoch",
            save_strategy="epoch",
            load_best_model_at_end=True,
            metric_for_best_model="eval_accuracy",
            greater_is_better=True,
        )
        
        # Custom metrics function
        def compute_metrics(eval_pred):
            predictions, labels = eval_pred
            predictions = np.argmax(predictions, axis=1)
            accuracy = accuracy_score(labels, predictions)
            return {"accuracy": accuracy}
        
        # Initialize trainer
        trainer = Trainer(
            model=self.model,
            args=training_args,
            train_dataset=train_dataset,
            eval_dataset=val_dataset,
            compute_metrics=compute_metrics,
        )
        
        # Train model
        logger.info("Starting model training...")
        train_result = trainer.train()
        
        # Evaluate model
        eval_result = trainer.evaluate(val_dataset)
        
        # Generate predictions for metrics
        predictions = trainer.predict(val_dataset)
        predicted_labels = np.argmax(predictions.predictions, axis=1)
        
        # Calculate detailed metrics
        class_report = classification_report(
            val_labels, predicted_labels, 
            target_names=unique_labels,
            output_dict=True
        )
        
        conf_matrix = confusion_matrix(val_labels, predicted_labels)
        
        # Create metrics object
        metrics = ModelMetrics(
            accuracy=eval_result["eval_accuracy"],
            precision=class_report["macro avg"]["precision"],
            recall=class_report["macro avg"]["recall"],
            f1_score=class_report["macro avg"]["f1-score"],
            class_metrics={
                label: {
                    "precision": class_report[label]["precision"],
                    "recall": class_report[label]["recall"],
                    "f1-score": class_report[label]["f1-score"]
                }
                for label in unique_labels
            },
            confusion_matrix=conf_matrix.tolist(),
            model_version="1.0",
            test_set_size=len(val_labels)
        )
        
        # Save model and metadata
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Save model and tokenizer
            trainer.save_model(output_dir)
            
            # Save class labels
            with open(output_path / "class_labels.json", 'w') as f:
                json.dump(unique_labels, f)
            
            # Save metrics
            with open(output_path / "metrics.json", 'w') as f:
                json.dump(metrics.dict(), f, indent=2, default=str)
            
            logger.info(f"Model saved to {output_dir}")
        
        # Update classifier pipeline
        self.classifier_pipeline = pipeline(
            "text-classification",
            model=self.model,
            tokenizer=self.tokenizer,
            device=0 if self.config.device == "cuda" else -1,
            return_all_scores=True
        )
        
        self.class_labels = unique_labels
        
        logger.info(f"Training completed. Accuracy: {metrics.accuracy:.3f}")
        return metrics
    
    def evaluate_model(self, test_data: List[Tuple[str, str]]) -> ModelMetrics:
        """
        Evaluate model performance on test data
        
        Args:
            test_data: List of (text, label) tuples for testing
            
        Returns:
            ModelMetrics with evaluation results
        """
        if not test_data:
            raise ValueError("Test data cannot be empty")
        
        texts, true_labels = zip(*test_data)
        
        # Get predictions
        predictions = []
        for text in texts:
            result = self.classify_document(
                DocumentContent(text=text, metadata={})
            )
            predictions.append(result.classification.document_type.value)
        
        # Calculate metrics
        accuracy = accuracy_score(true_labels, predictions)
        class_report = classification_report(
            true_labels, predictions, 
            output_dict=True
        )
        conf_matrix = confusion_matrix(true_labels, predictions)
        
        # Get unique labels
        unique_labels = list(set(true_labels + predictions))
        
        metrics = ModelMetrics(
            accuracy=accuracy,
            precision=class_report["macro avg"]["precision"],
            recall=class_report["macro avg"]["recall"],
            f1_score=class_report["macro avg"]["f1-score"],
            class_metrics={
                label: {
                    "precision": class_report.get(label, {}).get("precision", 0.0),
                    "recall": class_report.get(label, {}).get("recall", 0.0),
                    "f1-score": class_report.get(label, {}).get("f1-score", 0.0)
                }
                for label in unique_labels
            },
            confusion_matrix=conf_matrix.tolist(),
            model_version="1.0",
            test_set_size=len(test_data)
        )
        
        logger.info(f"Model evaluation completed. Accuracy: {accuracy:.3f}")
        return metrics
    
    def get_model_info(self) -> Dict[str, Any]:
        """Get information about the current model"""
        return {
            "model_name": self.config.model_name,
            "model_path": self.config.model_path,
            "num_labels": len(self.class_labels),
            "supported_categories": [cat.value for cat in self.supported_categories],
            "class_labels": self.class_labels,
            "device": self.config.device,
            "max_length": self.config.max_length,
            "confidence_threshold": self.config.confidence_threshold
        }